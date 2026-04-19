from flask import Flask, jsonify, request, send_from_directory
from datetime import datetime, timedelta
from io import BytesIO
import csv
import json
import os
import re
import sqlite3
import zipfile
from xml.etree import ElementTree as ET

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

from flask_cors import CORS
from openpyxl import load_workbook

app = Flask(__name__, static_folder='static', static_url_path='')
CORS(app)

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DB_FILE = os.path.join(ROOT_DIR, 'cases.json')
SQLITE_DB_FILE = os.path.join(ROOT_DIR, 'data', 'cases.db')
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

def load_db():
    if os.path.exists(DB_FILE):
        with open(DB_FILE, 'r') as f:
            db = json.load(f)
        return db
    return {'cases': [], 'next_id': 1, 'next_stt': 1}

def save_db(db):
    with open(DB_FILE, 'w') as f:
        json.dump(db, f, indent=2, default=str)

def parse_date_string(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    value = str(value).strip()
    # Handle ISO 8601 variants, including Zulu timezone and date-only strings
    try:
        if value.endswith('Z'):
            value = value[:-1] + '+00:00'
        return datetime.fromisoformat(value)
    except ValueError:
        pass

    parts = re.split(r'[\.\-/T ]', value)
    if len(parts) >= 3:
        try:
            day, month, year = [int(parts[i]) for i in range(3)]
            if year < 100:
                year += 2000
            return datetime(year, month, day)
        except ValueError:
            return None

    return None

def calculate_deadline(ngay_thu_ly, loai_an):
    date = parse_date_string(ngay_thu_ly)
    if not date:
        return ''
    if loai_an == "KDTM":
        return (date + timedelta(days=90)).isoformat()
    return (date + timedelta(days=180)).isoformat()


def parse_excel_date(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date().isoformat()
    if hasattr(value, 'to_pydatetime'):
        try:
            return value.to_pydatetime().date().isoformat()
        except Exception:
            pass
    value = str(value).strip()
    if not value:
        return None
    try:
        return datetime.fromisoformat(value).date().isoformat()
    except ValueError:
        pass
    parts = re.split(r'[\./\-\s]', value)
    if len(parts) == 3:
        try:
            day, month, year = [int(part) for part in parts]
            if year < 100:
                year += 2000
            return datetime(year, month, day).date().isoformat()
        except ValueError:
            return None
    return None


def normalize_bool(value):
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    return text in ['1', 'true', 'yes', 'x', 'đã mã hóa', 'da ma hoa', 'checked']


def read_txt_document(file_content: bytes) -> str:
    try:
        return file_content.decode('utf-8').strip()
    except UnicodeDecodeError:
        return file_content.decode('latin-1', errors='ignore').strip()


def read_csv_document(file_content: bytes) -> str:
    text = read_txt_document(file_content)
    rows = []
    for row in csv.reader(text.splitlines()):
        rows.append(' | '.join(row))
    return '\n'.join(rows)


FIELD_HEADER_ALIASES = {
    'stt': ['stt'],
    'so_thu_ly': ['số thủ lý', 'số thụ lý', 'sothuly', 'so_thu_ly', 'sothuly', 'sothuly', 'sốthuly'],
    'ngay_thu_ly': ['ngày thủ lý', 'ngày thụ lý', 'ngaythuly', 'ngay_thu_ly', 'ngaythu ly', 'ngay_thuly'],
    'duong_su': ['đương sự', 'duongsu', 'duong_su', 'duong su'],
    'quan_he_phap_luat': ['quan hệ pháp luật', 'quan he phap luat', 'quan_he_phap_luat', 'quan he phap luat', 'quanhephapluat'],
    'loai_an': ['loại án', 'loai_an', 'loai an', 'loai an'],
    'ngay_xet_xu': ['ngày xét xử', 'ngay xet xu', 'ngay_xet_xu', 'ngayxetxu'],
    'qd_cnstt': ['qđ cnstt', 'qd cnstt', 'qd_cnstt', 'qd_cnstt', 'qd cnstt'],
    'trang_thai_giai_quyet': ['trạng thái giải quyết', 'trang thai giai quyet', 'trang_thai_giai_quyet'],
    'ghi_chu': ['ghi chú', 'ghi chu', 'ghi_chu'],
    'ma_hoa': ['mã hóa', 'mã hoa', 'ma hoa', 'ma_hoa']
}

HEADER_CANONICAL_MAP = {}
for canonical, aliases in FIELD_HEADER_ALIASES.items():
    for alias in aliases:
        normalized_alias = re.sub(r'[^a-z0-9]', '', alias.lower())
        HEADER_CANONICAL_MAP[normalized_alias] = canonical


def normalize_header_token(header_value):
    if header_value is None:
        return ''
    return re.sub(r'[^a-z0-9]', '', str(header_value).strip().lower())


def canonicalize_header(header_value):
    normalized = normalize_header_token(header_value)
    return HEADER_CANONICAL_MAP.get(normalized, normalized)


def normalize_document_row(row):
    return {
        canonicalize_header(key): value
        for key, value in row.items()
        if key is not None
    }


def get_document_value(row, keys):
    normalized = normalize_document_row(row)
    for key in keys:
        canonical = canonicalize_header(key)
        if canonical in normalized and normalized[canonical] not in (None, ''):
            return normalized[canonical]
    return None


def parse_tabular_text(content: str) -> list[dict]:
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if len(lines) < 2:
        return []

    delimiter = None
    if '|' in lines[0]:
        delimiter = '|'
    elif '\t' in lines[0]:
        delimiter = '\t'
    elif ',' in lines[0]:
        delimiter = ','
    else:
        return []

    reader = csv.reader(lines, delimiter=delimiter)
    rows = list(reader)
    if len(rows) < 2:
        return []

    headers = [header.strip() for header in rows[0]]
    parsed = []
    for row in rows[1:]:
        row_dict = {
            headers[i]: row[i].strip() if i < len(row) else ''
            for i in range(len(headers))
        }
        if any(value for value in row_dict.values()):
            parsed.append(row_dict)
    return parsed


def parse_docx_table_rows(file_content: bytes) -> list[dict]:
    if DocxDocument is None:
        return []

    try:
        doc = DocxDocument(BytesIO(file_content))
    except Exception:
        return []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if not any(headers):
            continue

        parsed = []
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            if not any(values):
                continue
            row_dict = {
                headers[i]: values[i] if i < len(values) else ''
                for i in range(len(headers))
            }
            parsed.append(row_dict)
        if parsed:
            return parsed
    return []


def try_parse_document_rows(file_type: str, file_content: bytes, content: str) -> list[dict]:
    if file_type == 'docx':
        rows = parse_docx_table_rows(file_content)
        if rows:
            return rows
        return parse_tabular_text(content)
    if file_type in ('txt', 'csv'):
        return parse_tabular_text(content)
    return []


def build_case_from_document_row(db, row):
    normalized_row = normalize_document_row(row)
    stt_value = get_document_value(normalized_row, ['stt'])
    stt = db.get('next_stt', 1)
    if stt_value not in (None, '', float('nan')):
        try:
            stt = int(float(stt_value))
        except Exception:
            stt = db.get('next_stt', 1)

    case = {
        'id': db.get('next_id', 1),
        'stt': stt,
        'bien_lai_an_phi': str(get_document_value(normalized_row, ['bien_lai_an_phi', 'biên lai án phí', 'biên lai an phí']) or ''),
        'so_thu_ly': str(get_document_value(normalized_row, ['so_thu_ly', 'số thụ lý', 'số thủ lý']) or ''),
        'ngay_thu_ly': parse_excel_date(get_document_value(normalized_row, ['ngay_thu_ly', 'ngày thụ lý', 'ngày thủ lý'])),
        'duong_su': str(get_document_value(normalized_row, ['duong_su', 'đương sự']) or ''),
        'quan_he_phap_luat': str(get_document_value(normalized_row, ['quan_he_phap_luat', 'quan hệ pháp luật']) or ''),
        'loai_an': str(get_document_value(normalized_row, ['loai_an', 'loại án']) or ''),
        'ngay_xet_xu': parse_excel_date(get_document_value(normalized_row, ['ngay_xet_xu', 'ngày xét xử']) or ''),
        'qd_cnstt': str(get_document_value(normalized_row, ['qd_cnstt', 'qđ cnstt', 'qd cnstt']) or ''),
        'trang_thai_giai_quyet': str(get_document_value(normalized_row, ['trang_thai_giai_quyet', 'trạng thái giải quyết']) or 'Hòa giải thành'),
        'ghi_chu': str(get_document_value(normalized_row, ['ghi_chu', 'ghi chú']) or ''),
        'ma_hoa': normalize_bool(get_document_value(normalized_row, ['ma_hoa', 'mã hóa', 'mã hoa'])),
        'han_giai_quyet': ''
    }
    case['han_giai_quyet'] = calculate_deadline(case['ngay_thu_ly'], case.get('loai_an', ''))
    return case


def insert_case_from_document_row(db, row):
    case = build_case_from_document_row(db, row)
    db['cases'].append(case)
    db['next_id'] = db.get('next_id', 1) + 1
    db['next_stt'] = db.get('next_stt', 1) + 1
    insert_case_into_sqlite(case)
    return case


def import_document_rows(db, rows):
    imported = 0
    processed = 0
    skipped = 0
    errors = []

    for row_index, row in enumerate(rows, start=1):
        processed += 1
        if not any(value not in (None, '') for value in row.values()):
            skipped += 1
            continue
        try:
            insert_case_from_document_row(db, row)
            imported += 1
        except Exception as exc:
            skipped += 1
            errors.append({
                'row': row_index,
                'error': str(exc),
                'data': {k: str(v) for k, v in row.items()}
            })

    return imported, processed, skipped, errors


def read_docx_document(file_content: bytes) -> str:
    if DocxDocument is not None:
        doc = DocxDocument(BytesIO(file_content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    paragraphs.append(' | '.join(row_data))
        return '\n'.join(paragraphs)

    try:
        with zipfile.ZipFile(BytesIO(file_content)) as docx_zip:
            xml = docx_zip.read('word/document.xml')
    except zipfile.BadZipFile:
        raise ValueError('File .docx không hợp lệ')
    except KeyError:
        raise ValueError('Không tìm thấy word/document.xml trong file .docx')

    try:
        tree = ET.fromstring(xml)
    except ET.ParseError:
        raise ValueError('Không thể phân tích file .docx')

    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paragraphs = []
    for paragraph in tree.findall('.//w:p', namespace):
        texts = [node.text for node in paragraph.findall('.//w:t', namespace) if node.text]
        if texts:
            paragraphs.append(''.join(texts))

    if not paragraphs:
        raise ValueError('Không thể đọc nội dung từ file .docx')

    return '\n'.join(paragraphs)


def read_pdf_document(file_content: bytes) -> str:
    if PdfReader is None:
        raise ValueError('PyPDF2 chưa được cài đặt')
    pdf = PdfReader(BytesIO(file_content))
    pages = []
    for page in pdf.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text)
    return '\n---PAGE BREAK---\n'.join(pages)


def insert_case_into_sqlite(case):
    conn = sqlite3.connect(SQLITE_DB_FILE)
    try:
        conn.execute(
            '''CREATE TABLE IF NOT EXISTS cases (
                id INTEGER PRIMARY KEY,
                stt INTEGER,
                bien_lai_an_phi TEXT,
                so_thu_ly TEXT,
                ngay_thu_ly TEXT,
                duong_su TEXT,
                quan_he_phap_luat TEXT,
                loai_an TEXT,
                ngay_xet_xu TEXT,
                qd_cnstt TEXT,
                trang_thai_giai_quyet TEXT,
                ghi_chu TEXT,
                ma_hoa INTEGER,
                han_giai_quyet TEXT
            )'''
        )
        conn.execute(
            '''INSERT INTO cases (
                id, stt, bien_lai_an_phi, so_thu_ly, ngay_thu_ly,
                duong_su, quan_he_phap_luat, loai_an, ngay_xet_xu,
                qd_cnstt, trang_thai_giai_quyet, ghi_chu, ma_hoa,
                han_giai_quyet
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                case['id'], case['stt'], case['bien_lai_an_phi'], case['so_thu_ly'], case['ngay_thu_ly'],
                case['duong_su'], case['quan_he_phap_luat'], case.get('loai_an', ''), case['ngay_xet_xu'],
                case['qd_cnstt'], case['trang_thai_giai_quyet'], case['ghi_chu'], int(case['ma_hoa']),
                case['han_giai_quyet']
            )
        )
        conn.commit()
    finally:
        conn.close()

@app.get("/cases/")
def read_cases():
    db = load_db()
    return jsonify(db['cases'])

@app.get("/cases/<int:case_id>")
def read_case(case_id):
    db = load_db()
    for case in db['cases']:
        if case['id'] == case_id:
            return jsonify(case)
    return jsonify({"detail": "Case not found"}), 404

@app.post("/cases/")
def create_new_case():
    try:
        db = load_db()
        data = request.get_json(silent=True)
        if not isinstance(data, dict):
            return jsonify({"detail": "Yêu cầu phải gửi JSON hợp lệ."}), 400
        
        stt = data.get('stt')
        if stt is None:
            stt = db.get('next_stt', 1)
            db['next_stt'] = stt + 1
        
        case = {
            'id': db.get('next_id', 1),
            'stt': stt,
            'bien_lai_an_phi': data.get('bien_lai_an_phi', ''),
            'so_thu_ly': data.get('so_thu_ly', ''),
            'ngay_thu_ly': data.get('ngay_thu_ly'),
            'duong_su': data.get('duong_su', ''),
            'quan_he_phap_luat': data.get('quan_he_phap_luat', ''),
            'loai_an': data.get('loai_an', ''),
            'ngay_xet_xu': data.get('ngay_xet_xu'),
            'qd_cnstt': data.get('qd_cnstt', ''),
            'trang_thai_giai_quyet': data.get('trang_thai_giai_quyet', 'Hòa giải thành'),
            'ghi_chu': data.get('ghi_chu', ''),
            'ma_hoa': data.get('ma_hoa', False),
            'han_giai_quyet': calculate_deadline(data.get('ngay_thu_ly'), data.get('loai_an', ''))
        }
        
        db['cases'].append(case)
        db['next_id'] = db.get('next_id', 1) + 1
        save_db(db)
        
        return jsonify(case), 201
    except Exception as e:
        import traceback
        print(f"Error in create_new_case: {str(e)}")
        traceback.print_exc()
        return jsonify({"detail": f"Server error: {str(e)}"}), 500

@app.put("/cases/<int:case_id>")
def update_existing_case(case_id):
    db = load_db()
    data = request.get_json(silent=True)
    if not isinstance(data, dict):
        return jsonify({"detail": "Yêu cầu phải gửi JSON hợp lệ."}), 400
    
    for case in db['cases']:
        if case['id'] == case_id:
            if 'bien_lai_an_phi' in data:
                case['bien_lai_an_phi'] = data['bien_lai_an_phi']
            if 'so_thu_ly' in data:
                case['so_thu_ly'] = data['so_thu_ly']
            if 'ngay_thu_ly' in data:
                case['ngay_thu_ly'] = data['ngay_thu_ly']
            if 'duong_su' in data:
                case['duong_su'] = data['duong_su']
            if 'quan_he_phap_luat' in data:
                case['quan_he_phap_luat'] = data['quan_he_phap_luat']
            if 'ngay_xet_xu' in data:
                case['ngay_xet_xu'] = data['ngay_xet_xu']
            if 'qd_cnstt' in data:
                case['qd_cnstt'] = data['qd_cnstt']
            if 'trang_thai_giai_quyet' in data:
                case['trang_thai_giai_quyet'] = data['trang_thai_giai_quyet']
            if 'ghi_chu' in data:
                case['ghi_chu'] = data['ghi_chu']
            if 'ma_hoa' in data:
                case['ma_hoa'] = data['ma_hoa']
            if 'loai_an' in data:
                case['loai_an'] = data['loai_an']

            if 'ngay_thu_ly' in data or 'loai_an' in data:
                case['han_giai_quyet'] = calculate_deadline(case['ngay_thu_ly'], case.get('loai_an', ''))

            save_db(db)
            return jsonify(case)
    
    return jsonify({"detail": "Case not found"}), 404

@app.delete("/cases/<int:case_id>")
def delete_existing_case(case_id):
    db = load_db()

    for i, case in enumerate(db['cases']):
        if case['id'] == case_id:
            db['cases'].pop(i)
            for index, remaining_case in enumerate(db['cases'], start=1):
                remaining_case['stt'] = index
            db['next_stt'] = len(db['cases']) + 1
            save_db(db)
            return jsonify({"message": "Case deleted"})
    
    return jsonify({"detail": "Case not found"}), 404

@app.post('/upload-excel/')
def upload_excel():
    if 'file' not in request.files:
        return jsonify({"detail": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename.lower().endswith('.xlsx'):
        return jsonify({"detail": "File must be .xlsx"}), 400

    if pd is not None:
        try:
            df = pd.read_excel(file, engine='openpyxl')
        except Exception as e:
            return jsonify({"detail": f"Không thể đọc file Excel: {str(e)}"}), 400

        def get_value(row, keys):
            for key in keys:
                if key in row and pd.notna(row[key]):
                    return row[key]
            return None

        rows = []
        for _, row in df.iterrows():
            if row.isna().all():
                continue
            rows.append(row)
    else:
        workbook = load_workbook(file, data_only=True)
        sheet = workbook.active
        headers = [cell.value for cell in sheet[1]]

        def get_value(row_dict, keys):
            for key in keys:
                if key in row_dict and row_dict[key] is not None:
                    return row_dict[key]
            return None

        rows = []
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not any(cell is not None for cell in row):
                continue
            rows.append({headers[i]: row[i] for i in range(min(len(headers), len(row)))})

    db = load_db()
    added = 0

    for row in rows:
        ngay_thu_ly = parse_excel_date(get_value(row, ['Ngày Thụ Lý', 'Ngày thụ lý', 'ngày thụ lý']))
        if not ngay_thu_ly:
            continue

        ngay_xet_xu = parse_excel_date(get_value(row, ['Ngày Xét Xử', 'Ngày xét xử', 'ngày xét xử']))

        case = {
            'id': db.get('next_id', 1),
            'stt': db.get('next_stt', 1),
            'bien_lai_an_phi': str(get_value(row, ['Biên Lai Án Phí', 'Biên lai án phí', 'bien_lai_an_phi']) or ''),
            'so_thu_ly': str(get_value(row, ['Số Thụ Lý', 'Số thụ lý', 'so_thu_ly']) or ''),
            'ngay_thu_ly': ngay_thu_ly,
            'duong_su': str(get_value(row, ['Đương Sự', 'Tên đương sự', 'duong_su']) or ''),
            'quan_he_phap_luat': str(get_value(row, ['Quan Hệ Pháp Luật', 'Quan hệ pháp luật', 'quan_he_phap_luat']) or ''),
            'loai_an': str(get_value(row, ['Loại Án', 'loai_an']) or ''),
            'ngay_xet_xu': ngay_xet_xu,
            'qd_cnstt': str(get_value(row, ['QĐ Công Nhận STT', 'qd_cnstt']) or ''),
            'trang_thai_giai_quyet': str(get_value(row, ['Trạng Thái Giải Quyết', 'Trạng thái giải quyết', 'trang_thai_giai_quyet']) or 'Hòa giải thành'),
            'ghi_chu': str(get_value(row, ['Ghi Chú', 'ghi_chu']) or ''),
            'ma_hoa': normalize_bool(get_value(row, ['Đã Mã Hóa', 'Đã mã hóa', 'ma_hoa'])),
            'han_giai_quyet': ''
        }

        case['han_giai_quyet'] = calculate_deadline(case['ngay_thu_ly'], case.get('loai_an', ''))

        db['cases'].append(case)
        db['next_id'] = db.get('next_id', 1) + 1
        db['next_stt'] = db.get('next_stt', 1) + 1
        insert_case_into_sqlite(case)
        added += 1

    save_db(db)
    return jsonify({'message': f'Đã thêm {added} hồ sơ thành công'})

@app.post('/upload-document/')
def upload_document():
    if 'file' not in request.files:
        return jsonify({'detail': 'No file uploaded'}), 400

    file = request.files['file']
    filename_lower = file.filename.lower()
    file_content = file.read()

    upload_name = f"{datetime.now():%Y%m%d%H%M%S}_{os.path.basename(file.filename)}"
    upload_path = os.path.join(UPLOAD_DIR, upload_name)
    try:
        with open(upload_path, 'wb') as f:
            f.write(file_content)
    except Exception as exc:
        return jsonify({'detail': f'Không thể lưu file upload: {exc}'}), 500

    content = ''
    file_type = 'unknown'
    imported = 0
    processed = 0
    skipped = 0
    errors = []

    try:
        if filename_lower.endswith(('.xlsx', '.xls')):
            file_type = 'excel'
            if pd is not None:
                try:
                    df = pd.read_excel(BytesIO(file_content), engine='openpyxl')
                    headers = list(df.columns)
                    rows = []
                    for _, row in df.iterrows():
                        rows.append(' | '.join(str(row.get(col, '')) for col in headers))
                    content = '\n'.join([' | '.join(str(h) for h in headers)] + rows)
                except Exception as exc:
                    return jsonify({'detail': f'Không thể đọc file Excel: {exc}'}), 400
            else:
                try:
                    workbook = load_workbook(BytesIO(file_content), data_only=True)
                    sheet = workbook.active
                    headers = [cell.value for cell in sheet[1]] if sheet.max_row >= 1 else []
                    rows = []
                    for row in sheet.iter_rows(min_row=2, values_only=True):
                        if not any(cell is not None for cell in row):
                            continue
                        rows.append(' | '.join(str(cell) if cell is not None else '' for cell in row))
                    content = '\n'.join([' | '.join(str(h) for h in headers)] + rows)
                except Exception as exc:
                    return jsonify({'detail': f'Không thể đọc file Excel: {exc}'}), 400

        elif filename_lower.endswith('.docx'):
            file_type = 'docx'
            try:
                content = read_docx_document(file_content)
            except ValueError as exc:
                return jsonify({'detail': str(exc)}), 400
            except Exception as exc:
                return jsonify({'detail': f'Không thể đọc file Word (.docx): {exc}'}), 400

        elif filename_lower.endswith('.doc'):
            return jsonify({'detail': 'File .doc chưa được hỗ trợ. Vui lòng chuyển sang .docx hoặc .pdf'}), 400

        elif filename_lower.endswith('.pdf'):
            file_type = 'pdf'
            try:
                content = read_pdf_document(file_content)
            except ValueError as exc:
                return jsonify({'detail': str(exc)}), 400
            except Exception as exc:
                return jsonify({'detail': f'Không thể đọc file PDF: {exc}'}), 400

        elif filename_lower.endswith('.txt'):
            file_type = 'txt'
            content = read_txt_document(file_content)

        elif filename_lower.endswith('.csv'):
            file_type = 'csv'
            content = read_csv_document(file_content)

        else:
            return jsonify({'detail': 'Loại file không được hỗ trợ. Hỗ trợ: .xlsx, .xls, .docx, .pdf, .txt, .csv'}), 400

        rows = try_parse_document_rows(file_type, file_content, content)
        if rows:
            db = load_db()
            imported, processed, skipped, errors = import_document_rows(db, rows)
            save_db(db)
        else:
            imported = 0
            processed = 0
            skipped = 0
            errors = []

    except Exception as exc:
        return jsonify({'detail': str(exc)}), 500

    preview = content[:500] + '...' if len(content) > 500 else content
    payload = {
        'message': f'Đã đọc file {file_type.upper()} thành công',
        'file_type': file_type,
        'file_name': file.filename,
        'file_size': len(file_content),
        'content_length': len(content),
        'file_path': upload_path,
        'content': content[:1000] if len(content) > 1000 else content,
        'preview': preview,
        'imported': imported,
        'processed': processed,
        'skipped': skipped,
        'errors': errors,
    }

    if imported > 0:
        payload['message'] = f'Đã nhập {imported} hồ sơ từ file {file_type.upper()} thành công'

    return jsonify(payload)

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve_spa(path):
    if path != '' and os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    return send_from_directory(app.static_folder, 'index.html')

@app.get("/statistics/")
def get_stats():
    db = load_db()
    cases = db['cases']

    status_counts = {}
    overdue = 0
    warning = 0
    now = datetime.now()

    final_statuses = ["Hòa giải thành", "Xét xử", "Đình chỉ", "Tạm đình chỉ", "Bản án", "Kết thúc"]
    case_types = ["Dân sự", "Hôn nhân", "KDTM", "Lao động", "Hình sự", "Hành chính", "Cai nghiện"]
    type_counts = {}

    for case in cases:
        status = case.get('trang_thai_giai_quyet', 'Chưa cập nhật') or 'Chưa cập nhật'
        status_counts[status] = status_counts.get(status, 0) + 1

        case_type = case.get('loai_an', 'Khác') or 'Khác'
        type_counts[case_type] = type_counts.get(case_type, 0) + 1

        if status not in final_statuses:
            try:
                han_date = datetime.fromisoformat(case['han_giai_quyet'])
                if now > han_date:
                    overdue += 1
                elif (han_date - now).days < 15:
                    warning += 1
            except (ValueError, TypeError, KeyError):
                pass

    for status in final_statuses:
        if status not in status_counts:
            status_counts[status] = 0

    for case_type in case_types:
        if case_type not in type_counts:
            type_counts[case_type] = 0

    completed = sum(status_counts.get(s, 0) for s in final_statuses)
    ongoing = len(cases) - completed

    return jsonify({
        "status_counts": status_counts,
        "type_counts": type_counts,
        "dang_giai_quyet": ongoing,
        "an_sap_het_han": warning,
        "an_qua_han": overdue
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=False)
