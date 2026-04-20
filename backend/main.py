from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from backend.database import SessionLocal, initialize_database
from backend.crud import get_cases, get_case_by_id, create_case, update_case, delete_case, get_statistics
from backend.schemas import Case as CaseSchema, CaseCreate, CaseUpdate
from datetime import datetime
from io import BytesIO
import csv
import logging
import os
import tempfile
import unicodedata
from openpyxl import load_workbook

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import textract
except ImportError:
    textract = None

initialize_database()

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger('backend.upload')

app = FastAPI(title="Quản lý Hồ Sơ Án API")

static_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/dist"))

origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://localhost:5173",
    "http://localhost:5174",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:5174"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

SUPPORTED_DOCUMENT_EXTENSIONS = ('.xlsx', '.xls', '.docx', '.doc', '.pdf', '.txt', '.csv')

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.get("/cases/", response_model=list[CaseSchema])
def read_cases(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    cases = get_cases(db, skip=skip, limit=limit)
    return cases

@app.get("/cases/{case_id}", response_model=CaseSchema)
def read_case(case_id: int, db: Session = Depends(get_db)):
    db_case = get_case_by_id(db, case_id)
    if db_case is None:
        raise HTTPException(status_code=404, detail="Case not found")
    return db_case

@app.post("/cases/", response_model=CaseSchema)
def create_new_case(case: CaseCreate, db: Session = Depends(get_db)):
    return create_case(db, case)

@app.put("/cases/{case_id}", response_model=CaseSchema)
def update_existing_case(case_id: int, case: CaseUpdate, db: Session = Depends(get_db)):
    db_case = update_case(db, case_id, case)
    if db_case is None:
        raise HTTPException(status_code=404, detail="Case not found")
    return db_case

@app.delete("/cases/{case_id}")
def delete_existing_case(case_id: int, db: Session = Depends(get_db)):
    db_case = delete_case(db, case_id)
    if db_case is None:
        raise HTTPException(status_code=404, detail="Case not found")
    return {"message": "Case deleted"}

HEADER_FIELD_ALIASES = {
    'stt': {'stt', 'sothutu', 'so thutu', 'so thu tu', 'thutu', 'thu tu'},
    'bien_lai_an_phi': {
        'bienlaianphi', 'bien lai an phi', 'bien lai anphi', 'bien lai', 'bien lai an phi',
        'bienlai', 'bien lai an phi', 'bien laianphi'
    },
    'so_thu_ly': {'sothuly', 'so thu ly', 'so thu ly', 'sothuly'},
    'ngay_thu_ly': {'ngaythuly', 'ngay thu ly', 'ngay thuly', 'ngaythu ly', 'ngaythu ly'},
    'duong_su': {'tenduongsu', 'ten duong su', 'duong su', 'duong su', 'ten duong su', 'duongsu'},
    'quan_he_phap_luat': {
        'quanhephapluat', 'quan he phap luat', 'quan he phap luat', 'quan he phap luat',
        'quanhetranhchap', 'quan he tranh chap', 'lyhon', 'ly hon', 'tranhchap', 'tranh chap',
        'quanhe', 'quan he'
    },
    'loai_an': {'loaian', 'loai an', 'loaian', 'loai_an'},
    'ngay_xet_xu': {'ngayxetxu', 'ngay xet xu', 'ngayxet xu', 'ngay xetxu'},
    'qd_cnstt': {
        'qdcnstt', 'qd cnstt', 'qd cong nhan stt', 'qd cong nhan su thoa thuan',
        'qd cong nhan', 'qdcn', 'qd'
    },
    'trang_thai_giai_quyet': {
        'trangthai', 'trang thai', 'trang thai giai quyet', 'trangthaigiaiquyet',
        'trang thai giai quyet', 'trangthai giaiquyet'
    },
    'ghi_chu': {'ghichu', 'ghi chu', 'ghi chú', 'ghichu'},
    'ma_hoa': {'mahoa', 'mã hóa', 'da ma hoa', 'da ma hoa', 'ma hoa', 'mãhoa'}
}

CANONICAL_HEADER_MAP = {}


def normalize_header(value: str) -> str:
    if value is None:
        return ''
    text = str(value).strip().lower()
    text = unicodedata.normalize('NFD', text)
    text = ''.join(ch for ch in text if unicodedata.category(ch) != 'Mn')
    text = text.replace('đ', 'd')
    for ch in [' ', '_', '-', '.', ',', '(', ')', '/', '\\', '"', "'", ':']:
        text = text.replace(ch, '')
    return text


def get_canonical_header(value: str) -> str | None:
    normalized = normalize_header(value)
    return CANONICAL_HEADER_MAP.get(normalized)


for canonical, aliases in HEADER_FIELD_ALIASES.items():
    for alias in aliases:
        CANONICAL_HEADER_MAP[normalize_header(alias)] = canonical


def parse_excel_datetime(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    text = str(value).strip()
    if not text:
        return None
    if pd is not None:
        try:
            parsed = pd.to_datetime(text, dayfirst=True, errors='coerce')
            if parsed is not pd.NaT:
                return parsed.to_pydatetime()
        except Exception:
            pass
    for fmt in [
        '%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d', '%d.%m.%Y', '%d %m %Y',
        '%d/%m/%Y %H:%M:%S', '%Y-%m-%d %H:%M:%S', '%d-%m-%Y %H:%M:%S'
    ]:
        try:
            return datetime.strptime(text, fmt)
        except Exception:
            pass
    try:
        return datetime.fromisoformat(text)
    except Exception:
        return None


def parse_bool(value):
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    return text in ['1', 'true', 'yes', 'x', 'da ma hoa', 'da ma hoa', 'checked', 'y', 'co', 'có']


def _normalize_pandas_columns(columns):
    normalized = [normalize_header(col) for col in columns]
    return [get_canonical_header(col) or col for col in normalized]


def _choose_sheet(workbook):
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        if sheet.sheet_state == 'visible':
            return sheet
    return workbook.active


def _find_header_row(worksheet, max_rows=12, min_non_empty=2, min_known_fields=2):
    fallback = None
    for row_index in range(1, max_rows + 1):
        row = list(worksheet.iter_rows(min_row=row_index, max_row=row_index, values_only=True))[0]
        normalized = [normalize_header(cell) for cell in row]
        if sum(1 for value in normalized if value) >= min_non_empty and fallback is None:
            fallback = (row_index, row)
        known_fields = sum(1 for value in normalized if value in CANONICAL_HEADER_MAP)
        if known_fields >= min_known_fields:
            return row_index, row
    if fallback is not None:
        return fallback
    first_row = list(worksheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
    return 1, first_row


def read_excel_rows(file_content: bytes):
    rows = _read_excel_rows_openpyxl(file_content)
    if rows:
        logger.info(f"[read_excel_rows] openpyxl read {len(rows)} rows")
        return rows

    if pd is not None:
        try:
            df = pd.read_excel(BytesIO(file_content), engine='openpyxl', header=0)
            df.columns = _normalize_pandas_columns(df.columns)
            rows = [
                {col: row[col] for col in df.columns}
                for _, row in df.iterrows()
            ]
            logger.info(f"[read_excel_rows] pandas read {len(df)} rows, columns={list(df.columns)}")
            if rows:
                return rows
        except Exception as exc:
            logger.error(f"[read_excel_rows] pandas read failed: {exc}")

    return rows


def read_word_document(file_content: bytes, filename: str) -> str:
    """Đọc nội dung từ file Word (.docx, .doc)"""
    extension = os.path.splitext(filename)[1].lower()
    if extension == '.docx':
        if DocxDocument is None:
            raise HTTPException(status_code=400, detail="python-docx chưa được cài đặt")
        try:
            doc = DocxDocument(BytesIO(file_content))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        paragraphs.append(' | '.join(row_data))
            full_text = '\n'.join(paragraphs)
            logger.info(f"[read_word_document] read {len(paragraphs)} paragraphs/rows from DOCX")
            return full_text
        except Exception as exc:
            logger.error(f"[read_word_document] DOCX failed: {exc}")
            raise HTTPException(status_code=400, detail=f"Không thể đọc file Word (.docx): {exc}")

    if extension == '.doc':
        if textract is None:
            raise HTTPException(
                status_code=400,
                detail="File .doc chưa được hỗ trợ nếu không có textract. Hãy chuyển đổi sang .docx hoặc cài textract và các công cụ chuyển đổi tương ứng."
            )
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            text = textract.process(temp_path)
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='ignore')
            content = text.strip()
            logger.info(f"[read_word_document] read {len(content.splitlines())} lines from DOC")
            return content
        except Exception as exc:
            logger.error(f"[read_word_document] DOC failed: {exc}")
            raise HTTPException(status_code=400, detail=f"Không thể đọc file Word (.doc): {exc}")
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass

    raise HTTPException(status_code=400, detail="Định dạng Word không được hỗ trợ")


def read_txt_document(file_content: bytes) -> str:
    try:
        text = file_content.decode('utf-8')
    except UnicodeDecodeError:
        text = file_content.decode('latin-1', errors='ignore')
    return text.strip()


def read_csv_document(file_content: bytes) -> str:
    try:
        text = file_content.decode('utf-8')
    except UnicodeDecodeError:
        text = file_content.decode('latin-1', errors='ignore')
    rows = []
    for row in csv.reader(text.splitlines()):
        rows.append(' | '.join(row))
    return '\n'.join(rows)


def read_pdf_document(file_content: bytes) -> str:
    """Đọc nội dung từ file PDF"""
    if PdfReader is None:
        raise HTTPException(status_code=400, detail="PyPDF2 chưa được cài đặt")
    
    try:
        pdf = PdfReader(BytesIO(file_content))
        pages = []
        
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
        
        full_text = '\n---PAGE BREAK---\n'.join(pages)
        logger.info(f"[read_pdf_document] read {len(pdf.pages)} pages")
        return full_text
    except Exception as exc:
        logger.error(f"[read_pdf_document] failed: {exc}")
        raise HTTPException(status_code=400, detail=f"Không thể đọc file PDF: {exc}")


def _build_headers(header_row: list) -> list[str]:
    normalized = [normalize_header(value) for value in header_row]
    headers = []
    seen = {}
    for index, header in enumerate(normalized):
        if not header:
            header_name = f'col{index + 1}'
        else:
            header_name = get_canonical_header(header) or header
        count = seen.get(header_name, 0)
        seen[header_name] = count + 1
        if count:
            header_name = f'{header_name}_{count + 1}'
        headers.append(header_name)
    return headers


def _read_excel_rows_openpyxl(file_content: bytes):
    try:
        workbook = load_workbook(BytesIO(file_content), data_only=True)
        worksheet = _choose_sheet(workbook)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Không thể đọc file Excel: {exc}")

    start_row, header_row = _find_header_row(worksheet)
    headers = _build_headers(header_row)
    if not any(headers):
        return []

    rows = []
    blank_count = 0
    for row in worksheet.iter_rows(min_row=start_row + 1, values_only=True):
        if row is None or all(cell is None or str(cell).strip() == '' for cell in row):
            blank_count += 1
            if blank_count >= 10:
                break
            continue
        blank_count = 0
        row_data = {
            headers[idx]: row[idx]
            for idx in range(min(len(headers), len(row)))
            if headers[idx]
        }
        if not _has_row_data(row_data):
            continue
        rows.append(row_data)
    return rows


def _is_blank_value(value):
    if value is None:
        return True
    if isinstance(value, str) and not value.strip():
        return True
    try:
        if pd is not None and pd.isna(value):
            return True
    except Exception:
        pass
    try:
        return value != value
    except Exception:
        return False


def _has_row_data(row_data: dict):
    return any(not _is_blank_value(value) for value in row_data.values())


def get_excel_value(row_data: dict, *keys):
    for key in keys:
        normalized = normalize_header(key)
        canonical = get_canonical_header(normalized)
        if canonical and canonical in row_data:
            return row_data[canonical]
        if key in row_data:
            return row_data[key]
    return None


@app.post("/upload-excel/")
def upload_excel(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename.lower().endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="File must be .xlsx or .xls")

    file_content = file.file.read()
    upload_name = f"{datetime.now():%Y%m%d%H%M%S}_{os.path.basename(file.filename)}"
    upload_path = os.path.join(UPLOAD_DIR, upload_name)
    try:
        with open(upload_path, 'wb') as upload_file:
            upload_file.write(file_content)
        logger.info(f"[upload_excel] saved uploaded file to {upload_path} ({len(file_content)} bytes)")
    except Exception as exc:
        logger.error(f"[upload_excel] failed to save uploaded file: {exc}")
        raise HTTPException(status_code=500, detail=f"Không thể lưu file upload: {exc}")

    rows = read_excel_rows(file_content)
    if not rows:
        logger.info(f"[upload_excel] no rows found in uploaded file: {upload_path}")
        return {"message": "File Excel không có dữ liệu", "imported": 0, "processed": 0, "skipped": 0, "errors": [], "file_path": upload_path}

    imported = 0
    processed = 0
    skipped = 0
    errors = []

    for row_index, row_data in enumerate(rows, start=2):
        if not _has_row_data(row_data):
            skipped += 1
            continue

        processed += 1
        stt_value = get_excel_value(row_data, 'STT', 'so thu ly', 'stt')
        stt = None
        try:
            if stt_value not in (None, '', float('nan')):
                stt = int(float(stt_value))
        except Exception:
            stt = None

        case_data = CaseCreate(
            stt=stt,
            bien_lai_an_phi=str(get_excel_value(row_data, 'Biên lai án phí', 'bien lai an phi') or '').strip(),
            so_thu_ly=str(get_excel_value(row_data, 'Số thụ lý', 'so thu ly') or '').strip(),
            ngay_thu_ly=parse_excel_datetime(get_excel_value(row_data, 'Ngày thụ lý', 'ngay thu ly')),
            duong_su=str(get_excel_value(row_data, 'Tên đương sự', 'ten duong su', 'duong su') or '').strip(),
            quan_he_phap_luat=str(get_excel_value(row_data, 'Quan hệ tranh chấp', 'quan he tranh chap') or '').strip(),
            loai_an=str(get_excel_value(row_data, 'Loại án', 'loai an') or '').strip(),
            ngay_xet_xu=parse_excel_datetime(get_excel_value(row_data, 'Ngày xét xử', 'ngay xet xu')),
            qd_cnstt=str(get_excel_value(row_data, 'QĐ CNSTT', 'qd_cnstt', 'qd cnstt') or '').strip(),
            trang_thai_giai_quyet=str(get_excel_value(row_data, 'Trạng thái', 'trang thai') or 'Hòa giải thành').strip(),
            ghi_chu=str(get_excel_value(row_data, 'Ghi chú', 'ghi chu') or '').strip(),
            ma_hoa=parse_bool(get_excel_value(row_data, 'Mã hóa', 'ma hoa', 'ma_hoa'))
        )

        try:
            create_case(db, case_data)
            imported += 1
        except Exception as exc:
            db.rollback()
            logger.error(f"[upload_excel] create case failed row {row_index}: {exc}")
            errors.append({
                "row": row_index,
                "error": str(exc),
                "data": {k: str(v) for k, v in row_data.items()}
            })

    logger.info(f"[upload_excel] commit summary uploaded file {upload_path}: processed={processed}, imported={imported}, skipped={skipped}, errors={len(errors)}")
    result = {
        "message": f"Đã nhập {imported}/{processed} dòng đã xử lý.",
        "imported": imported,
        "processed": processed,
        "skipped": skipped,
        "errors": errors,
        "headers": list(rows[0].keys()),
        "file_path": upload_path
    }
    if errors:
        result["message"] += f" {len(errors)} dòng lỗi."
    return result


@app.post("/upload-document/")
def upload_document(file: UploadFile = File(...)):
    """
    Đọc nội dung từ các loại file khác nhau (Excel, Word, PDF)
    """
    filename_lower = file.filename.lower()
    file_content = file.file.read()
    
    upload_name = f"{datetime.now():%Y%m%d%H%M%S}_{os.path.basename(file.filename)}"
    upload_path = os.path.join(UPLOAD_DIR, upload_name)
    
    try:
        with open(upload_path, 'wb') as upload_file:
            upload_file.write(file_content)
        logger.info(f"[upload_document] saved file to {upload_path} ({len(file_content)} bytes)")
    except Exception as exc:
        logger.error(f"[upload_document] failed to save file: {exc}")
        raise HTTPException(status_code=500, detail=f"Không thể lưu file: {exc}")
    
    content = ""
    file_type = "unknown"
    
    # Xác định loại file và đọc nội dung tương ứng
    if filename_lower.endswith(('.xlsx', '.xls')):
        file_type = "excel"
        try:
            rows = read_excel_rows(file_content)
            if rows:
                headers = list(rows[0].keys())
                content_lines = [' | '.join(str(h) for h in headers)]
                for row in rows:
                    row_values = [str(row.get(h, '')) for h in headers]
                    content_lines.append(' | '.join(row_values))
                content = '\n'.join(content_lines)
                logger.info(f"[upload_document] read {len(rows)} rows from Excel")
            else:
                content = "File Excel không có dữ liệu"
        except Exception as exc:
            logger.error(f"[upload_document] excel read failed: {exc}")
            raise HTTPException(status_code=400, detail=f"Không thể đọc file Excel: {exc}")
    
    elif filename_lower.endswith(('.docx', '.doc')):
        file_type = "word"
        content = read_word_document(file_content, file.filename)

    elif filename_lower.endswith('.txt'):
        file_type = "text"
        content = read_txt_document(file_content)

    elif filename_lower.endswith('.csv'):
        file_type = "csv"
        content = read_csv_document(file_content)

    elif filename_lower.endswith('.pdf'):
        file_type = "pdf"
        content = read_pdf_document(file_content)
    
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Loại file không được hỗ trợ. Hỗ trợ: {', '.join(SUPPORTED_DOCUMENT_EXTENSIONS)}"
        )
    
    return {
        "message": f"Đã đọc file {file_type.upper()} thành công",
        "file_type": file_type,
        "file_name": file.filename,
        "file_path": upload_path,
        "file_size": len(file_content),
        "content": content[:1000] if len(content) > 1000 else content,
        "content_length": len(content),
        "preview": content[:500] + "..." if len(content) > 500 else content
    }

@app.get("/statistics/")
def get_stats(db: Session = Depends(get_db)):
    return get_statistics(db)

if os.path.isdir(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="frontend")


if __name__ == '__main__':
    import uvicorn

    port = int(os.environ.get('PORT', 8001))
    uvicorn.run(app, host='0.0.0.0', port=port)
